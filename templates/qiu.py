import datetime
from itertools import product

from cytoolz.curried import curry  # type:ignore
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pymonad.reader import Pipe  # type:ignore

from .base import (
    Color,
    Column,
    FontSize,
    color,
    font_size,
    init_blank_document,
    load_data,
)


@curry
def beginning(doc: Document) -> Document:
    doc.add_paragraph("{填写标题}", style="Title")

    paragraph = doc.add_paragraph()
    paragraph.add_run("出差时间：").bold = True
    paragraph.add_run("{填写起始日期}至{填写结束日期}")

    paragraph = doc.add_paragraph()
    paragraph.add_run("出差地点：").bold = True
    paragraph.add_run("{填写出差地点}")

    paragraph = doc.add_paragraph()
    paragraph.add_run("报销时间：").bold = True
    paragraph.add_run("{填写报销时间}")

    doc.add_paragraph("-" * 50)
    paragraph = doc.add_paragraph("出差前预支款：")
    paragraph.add_run("{填写预支款}")
    doc.add_paragraph("-" * 50)

    return doc


@curry
def summary(doc: Document, csv_file: str) -> Document:
    paragraph = doc.add_paragraph()
    paragraph.add_run("总花销费用和最终转入账户费用：").bold = True

    df = load_data(csv_file)
    categories = list(df[Column.CATEGORY].unique())
    names = list(df[Column.NAME].unique())
    rows = len(categories) + 2
    columns = len(names) + 2
    table = doc.add_table(rows=rows, cols=columns, style="Table Grid")
    for i in range(len(names)):
        table.cell(0, i + 1).text = names[i]
    table.cell(0, columns - 1).text = "合计"
    for i in range(len(categories)):
        table.cell(i + 1, 0).text = categories[i]
    table.cell(rows - 1, 0).text = "合计"

    def func(i: int, j: int) -> str:
        category = categories[i]
        name = names[j]
        amount = round(
            df[(df[Column.CATEGORY] == category) & (df[Column.NAME] == name)][
                Column.AMOUNT
            ].sum(),
            2,
        )
        return str(amount)

    for i, j in product(range(len(categories)), range(len(names))):
        table.cell(i + 1, j + 1).text = func(i, j)
    for i in range(len(categories)):
        table.cell(i + 1, columns - 1).text = str(
            round(df[df[Column.CATEGORY] == categories[i]][Column.AMOUNT].sum(), 2)
        )
    for i in range(len(names)):
        table.cell(rows - 1, i + 1).text = str(
            round(df[df[Column.NAME] == names[i]][Column.AMOUNT].sum(), 2)
        )
    table.cell(rows - 1, columns - 1).text = str(round(df[Column.AMOUNT].sum(), 2))
    doc.add_paragraph()

    paragraph = doc.add_paragraph()
    paragraph.add_run("无发票总费用：").bold = True

    paragraph = doc.add_paragraph()
    paragraph.add_run("最终决算：").bold = True
    doc.add_paragraph("{填写决算}")

    doc.add_page_break()

    return doc


@curry
def expense_statement(doc: Document, csv_file: str) -> Document:
    def func(x: int) -> Column:
        match x:
            case 0:
                return Column.NAME
            case 1:
                return Column.DATE
            case 2:
                return Column.DESCRIPTION
            case 3:
                return Column.AMOUNT
            case _:
                raise ValueError(f"Invalid index: {x}.")

    doc.add_heading(text="费用明细", level=1)

    headings = ("姓名", "日期", "事项", "费用", "合计")
    for i, (category, df) in enumerate(load_data(csv_file).groupby(Column.CATEGORY)):
        doc.add_paragraph(f"（{i + 1}）{category}费用：")
        paragraph_summary = doc.add_paragraph()
        paragraph_summary.add_run("支付组成：").bold = True

        heading_rows = 1
        rows = df.shape[0] + heading_rows
        table = doc.add_table(rows=rows, cols=len(headings), style="Table Grid")
        # Fill headings
        for i, heading in enumerate(headings):
            table.cell(0, i).text = heading
        # Fill isolate data
        for i, j in product(range(heading_rows, rows), range(1, 4)):
            table.cell(i, j).text = str(df[func(j)].iloc[i - heading_rows])
        # todo: Fill summary data
        i = heading_rows
        for name, df in df.groupby(Column.NAME):
            cell_name = table.cell(i, 0)
            cell_name.text = name

            cell_sum = table.cell(i, len(headings) - 1)
            amount_sum = round(df[Column.AMOUNT].sum(), 2)
            cell_sum.text = str(amount_sum)

            i += df.shape[0]
            cell_name.merge(table.cell(i - 1, 0))
            cell_sum.merge(table.cell(i - 1, len(headings) - 1))

            run = paragraph_summary.add_run(f"{name} ")
            run.bold = True
            run = paragraph_summary.add_run(f"￥{amount_sum} ")
            run.bold = True
            run.font.color.rgb = color(Color.BLUE)

        doc.add_paragraph()

    return doc


@curry
def confirm(doc: Document, csv_file: str) -> Document:
    names = list(load_data(csv_file).groupby(Column.NAME).groups.keys())

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"以上支付明细详单\r经出差{len(names)}人确认无误！")
    run.style = "Strong"  # type:ignore
    run.font.size = font_size(FontSize.一号)
    run.font.color.rgb = color(Color.RED)

    doc.add_paragraph()

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(str(" ".join(names)))
    run.style = "Emphasis"  # type:ignore
    run.bold = True

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(f"{datetime.datetime.now().strftime('%Y年%m月%d日')}")
    run.font.bold = True

    return doc


def main(csv_file: str) -> Document:
    return (
        Pipe(init_blank_document())
        .map(beginning)
        .map(summary(csv_file=csv_file))
        .map(expense_statement(csv_file=csv_file))
        .map(confirm(csv_file=csv_file))
        .flush()
    )
